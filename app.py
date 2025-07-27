# app.py

import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# Configure Gemini API
genai.configure(api_key="AIzaSyApYW1IwkubnWt77vY3kXm_USry03CY3Gc")
model = genai.GenerativeModel(model_name="models/gemini-1.5-flash")

st.set_page_config(page_title="AI-Powered Court Hearing Tool", layout="wide")
st.title("🧠 EAIEPL - AI Powered Court Hearing Preparation Tool")

# Sidebar Inputs
st.sidebar.header("Case Information")
case_name = st.sidebar.text_input("Case Name")
case_type = st.sidebar.selectbox("Case Type", ["Civil", "Criminal", "Other"])
facts = st.sidebar.text_area("Facts of the Case", height=200)
evidence = st.sidebar.text_area("Key Evidence", height=120)

# Advanced Inputs (on main page)
st.markdown("### 👨‍⚖️ Additional Case Details")
lawyer_name = st.text_input("Your Name (Lawyer)")
your_client = st.text_input("Your Client Name")
opposing_party = st.text_input("Opposing Party Name")
side = st.selectbox("You are representing", ["Plaintiff", "Defendant"])

st.markdown("### ⚙️ Generate Court Preparation Content")

# Output Buttons
col1, col2, col3 = st.columns(3)
generate_summary = col1.button("📄 Generate Summary")
generate_arguments = col2.button("⚖️ Generate Legal Arguments")
generate_questions = col3.button("❓ Generate Legal Questions")

col4, col5, col6 = st.columns(3)
generate_witness = col4.button("👥 Witness & Evidence")
generate_examination = col5.button("🧾 Examination & Cross")
generate_counter = col6.button("🛡️ Opponent Counter Questions")

# NEW Buttons
col7, col8, col9 = st.columns(3)
generate_full_report = col7.button("📥 Download Full Report")
find_similar_cases = col8.button("🔍 Find Similar Cases")
find_laws_sections = col9.button("📚 Find Applicable Laws & Sections")

# Init session state
for key in ['summary', 'arguments', 'questions', 'witness', 'examination', 'counter', 'similar_cases', 'law_sections']:
    if key not in st.session_state:
        st.session_state[key] = None

# Helper - Gemini generate
def generate_content(prompt):
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"❌ Error: {e}"

# Helper - Word Download
def generate_docx(title, content):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# 1. Summary
if generate_summary and facts.strip():
    prompt = f"""You are a legal assistant based in India. Generate a concise courtroom summary based on Indian legal standards using proper legal English.

Case Name: {case_name}
Case Type: {case_type}
Lawyer: {lawyer_name}
Client: {your_client}
Opposition: {opposing_party}
Side Represented: {side}
Facts: {facts}
Evidence: {evidence}

Generate Summary:"""
    st.session_state.summary = generate_content(prompt)

# 2. Legal Arguments
if generate_arguments and facts.strip():
    prompt = f"""You are an Indian legal analyst. Generate formal legal arguments for court submission based on Indian Constitution, IPC, Evidence Act, and relevant case laws.

Use proper citations and courtroom structure.

Case Name: {case_name}
Case Type: {case_type}
Lawyer: {lawyer_name}
Client: {your_client}
Side: {side}
Facts: {facts}
Evidence: {evidence}

Generate Legal Arguments:"""
    st.session_state.arguments = generate_content(prompt)

# 3. Legal Questions
if generate_questions and facts.strip():
    prompt = f"""You are assisting an Indian courtroom hearing. Frame legal questions in English for witness examination and cross-examination as per Indian court procedure.

Case Name: {case_name}
Client: {your_client}
Opposition: {opposing_party}
Facts: {facts}
Evidence: {evidence}

Generate formal courtroom-style legal questions:"""
    st.session_state.questions = generate_content(prompt)

# 4. Witness & Evidence
if generate_witness and facts.strip():
    prompt = f"""Act as a legal prep assistant in India. Based on the facts and evidence, prepare a structured breakdown of:

- Key Witnesses
- Role of Each Witness
- Primary Evidence & Its Relevance
- Legal Relevance under Indian Evidence Act

Case: {case_name}
Client: {your_client}
Opposition: {opposing_party}
Side: {side}
Facts: {facts}
Evidence: {evidence}

Generate Witness & Evidence Summary:"""
    st.session_state.witness = generate_content(prompt)

# 5. Examination & Cross
if generate_examination and facts.strip():
    prompt = f"""You are preparing an Indian lawyer for courtroom strategy.

Generate:
- Examination-in-Chief Questions (from {lawyer_name} to {your_client})
- Cross-Examination Questions (to {opposing_party}'s witnesses)

Follow courtroom tone and Indian legal standards.

Case: {case_name}
Side Represented: {side}
Facts: {facts}
Evidence: {evidence}
"""
    st.session_state.examination = generate_content(prompt)

# 6. Counter Questions
if generate_counter and facts.strip():
    prompt = f"""Act as the opposing counsel in this Indian court case.

Generate possible tough counter questions that could be asked to {your_client} during cross-examination.

Use Indian legal reasoning and courtroom English.

Facts: {facts}
Evidence: {evidence}
"""
    st.session_state.counter = generate_content(prompt)

# 7. Similar Case Finder
if find_similar_cases and facts.strip():
    prompt = f"""
You are a legal research assistant in India. Based on the given facts and case type, find 2–3 similar Indian court cases (Supreme or High Court preferred). Include short summaries and citations.

Case Type: {case_type}
Facts: {facts}
    """
    st.session_state.similar_cases = generate_content(prompt)

# 8. Applicable Laws & Sections
if find_laws_sections and facts.strip():
    prompt = f"""
You are an expert in Indian law. Based on the facts and case type, identify applicable Indian laws, statutes, and sections (IPC, CrPC, Evidence Act, Constitution, etc.) relevant to the case. Provide section numbers with brief explanations.

Case Type: {case_type}
Facts: {facts}
    """
    st.session_state.law_sections = generate_content(prompt)

# 9. Entire Report Generation
if generate_full_report:
    full_doc = Document()
    full_doc.add_heading("Court Preparation Report", level=1)
    for section_title, key in [
        ("Summary", "summary"),
        ("Legal Arguments", "arguments"),
        ("Legal Questions", "questions"),
        ("Witness & Evidence", "witness"),
        ("Examination & Cross Examination", "examination"),
        ("Opponent Counter Questions", "counter"),
        ("Similar Cases", "similar_cases"),
        ("Applicable Laws & Sections", "law_sections")
    ]:
        if st.session_state[key]:
            full_doc.add_heading(section_title, level=2)
            full_doc.add_paragraph(st.session_state[key])
    buffer = BytesIO()
    full_doc.save(buffer)
    buffer.seek(0)
    st.download_button("📥 Download Full Court Report", buffer, file_name=f"full_court_report.docx")

# Display Outputs

def render_output(title, content):
    if content:
        st.subheader(f"🧾 {title}")
        st.text_area(f"{title} Output", content, height=300)

render_output("📄 Summary", st.session_state.summary)
render_output("⚖️ Legal Arguments", st.session_state.arguments)
render_output("❓ Legal Questions", st.session_state.questions)
render_output("👥 Witness & Evidence", st.session_state.witness)
render_output("🧾 Examination & Cross Examination", st.session_state.examination)
render_output("🛡️ Opponent Counter Questions", st.session_state.counter)
render_output("🔍 Similar Cases", st.session_state.similar_cases)
render_output("📚 Applicable Laws & Sections", st.session_state.law_sections)
